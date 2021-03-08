#!/usr/bin/env python3
#-*- coding: utf-8 -*-

"""
分析 word 内容，提取题目、选项
"""

import os
import sys
import docx
import re
import json
import xlwt
import xlrd
from xlutils.copy import copy

xlsx_name = "消防安全责任人、负责人试卷.xlsx"
doc_name = "消防安全责任人、负责人试卷.docx"

curr_dir = os.path.dirname(os.path.realpath(__file__))
docx_file = os.path.join(curr_dir, doc_name)
xlsx_file = os.path.join(curr_dir, xlsx_name)



def extract_type(text):
    """
    从字符串中提取出单选题、多选题、判断题的字样
    """
    regex = "(单选题|多选题|判断题)"
    return re.search(regex, text).group(0)
    
def extract_decide(text):
    """
    提取判断题
    """
    regex = "(√|×)"
    result = re.search(regex, text)
    if result is None:
        return (None, None)
    answer = result.group(0)
    text = re.sub(regex, "", text)
    text = re.sub("^\s?\d+\.\s?", "", text).rstrip()
    return (text, answer)
    
def extract_choice(text):
    """
    提取选择题
    """
    regex = "([（\(]\s*)([A-Z]+)(\s*[）\)])"
    result = re.search(regex, text)
    if result is None:
        return (None, None)
    answer = result.group(2)
    text = re.sub(regex, r"\1\3", text)
    text = re.sub("^\s*\d+\.\s*", "", text).rstrip()
    return (text, answer)

def isMultipleSelector(text):
    """
    文本是否包含多个选项
    """
    regex = r"\(?(A|B|C|D|E|F|G|H)\)?\s?\w+"
    results = re.findall(regex, text) 
    if len(results) > 2:
        return True

    regex = r"\(?(A|B|C|D|E|F|G|H)\)?\s?[\w\u4e00-\u9fa5]+"
    results = re.findall(regex, text) 
    return len(results) > 2
    
def extractMultipleSelector(text):
    regex = r"\(?[A|B|C|D|E|F|G|H]\)?\s?([\s,\u3002\uff1b\uff0c\uff1a\u201c\u201d\uff08\uff09\u3001\uff1f\u300a\u300b\w]+)"
    results = re.findall(regex, text) 
    if len(results) > 2:
        return results

    regex = "\(?[A|B|C|D|E|F|G|H]\)\s?([\w\s,\u3002\uff1b\uff0c\uff1a\u201c\u201d\uff08\uff09\u3001\uff1f\u300a\u300b|\u4e00-\u9fa5]+)"
    results = re.findall(regex, text)
    return results
    
def remove_number(text):
    text = re.sub("^\s*\d+\.\s*", "", text).rstrip()
    return text
    
def isNumberStart(text):
    search = re.match("^\s*\d+\.\s*", text)
    if search is not None:
        return True
    
    search = re.match("^\s*\(?(A|B|C|D|E|F|G|H)", text) 
    print(text, search is None)
    if search is None:
        return True

    return False

def export_excel(boxs):
    """
    docx中已有答案提取导出
    """
    excel = xlwt.Workbook()
    sheet = excel.add_sheet("Sheet1") 
    headers = [
        "序号",	"考题标题", "考题答案", 	
        "选项1", "选项2", "选项3",  "选项4", "选项5"
    ]
    
    # 写标题行
    row_header = sheet.row(0)
    for i in range(len(headers)):
        row_header.write(i, headers[i]) 

    number = 1 # 序号
    for question in boxs.items():
        row = sheet.row(number)
        selects = question.get("selector", [])
        if len(selects) < 5:
            selects += [ "" for x in range(0, 5 - len(selects))]
        item = [
            number,
            question["question"], # 考题标题
            # questype, # 考题类型 
            question["answer"], # 考题答案
            *selects, # 5个选项
        ]
        for cell in range(len(item)):
            row.write(cell, item[cell])
        
        number += 1
            
    filename = "测试试卷.xlsx"
    excel.save(filename)

def write_excel_xls_append(boxs):
    """
    xlsx 中有答案提取追加
    """

    workbook = xlrd.open_workbook(xlsx_name)  # 打开工作簿
    sheets = workbook.sheet_names()  # 获取工作簿中的所有表格
    worksheet = workbook.sheet_by_name(sheets[0])  # 获取工作簿中所有表格中的的第一个表格
    rows_old = worksheet.nrows  # 获取表格中已存在的数据的行数
    new_workbook = copy(workbook)  # 将xlrd对象拷贝转化为xlwt对象
    new_worksheet = new_workbook.get_sheet(0)  # 获取转化后工作簿中的第一个表格

    # for i in range(0, index):
    #     for j in range(0, len(value[i])):
    #         new_worksheet.write(i+rows_old, j, value[i][j])  # 追加写入数据，注意是从i+rows_old行开始写入

    number = 1 # 序号
    for question in boxs:
    # for question in questions:
        row = new_worksheet.row(number)
        selects = question.get("selector", [])
        if len(selects) < 5:
            selects += [ "" for x in range(0, 5 - len(selects))]
        item = [
            number,
            question["question"], # 考题标题
            # questype, # 考题类型 
            question["answer"], # 考题答案
            *selects, # 5个选项
        ]
        for cell in range(len(item)):
            if cell == 2:
                continue
            row.write(cell, item[cell])
        
        number += 1

    new_workbook.save(xlsx_name)  # 保存工作簿
    print("xls格式表格【追加】写入数据成功！")

def analyze_document():
    doc = docx.Document(docx_file)
    boxs = []
    # {
    #     "单选题": [],
    #     "多选题": [],
    #     "判断题": [],
    # }
    types = ["单选题", "多选题", "判断题"]
    
    questype = ""
    selectorStart = False # 是否开始提取选项
    selector = []
    item = None
    
    with open("log.txt", "w", encoding="utf-8") as f:
        data = []
        for paragraph in doc.paragraphs:
            d = [paragraph.style.name, paragraph.text ]
            data.append(d) 
        f.write(json.dumps(data,indent=4, ensure_ascii=False))
    
    echo = False
    for paragraph in doc.paragraphs:
        # 提取题目类型 
        styleName = paragraph.style.name     
        if styleName == "Heading 2":
            type = extract_type(paragraph.text) 
            if type in types:
                questype = type

        text = paragraph.text 
        if styleName != "Normal" or questype == "" or text == "":
            continue

        question = answer = None

        
        if isNumberStart(text) and selectorStart: # 上一道题选项提取完毕
            item["selector"] = selector
            boxs.append(item)
            # 重置状态
            selectorStart = False
            item = None
            selector = []

        if questype == "判断题":
            # 有答案
            # [ question, answer ] = extract_decide(text)
            # if question is None or answer is None:
            #     continue
            
            # 没有答案
            question = remove_number(text) 
            answer = ""
            item = {
                "question": question,
                "answer": answer,
            }
            boxs.append(item)
            continue

        # 提取单选题
        if questype == "单选题" or questype == "多选题":
            # if len(selector) > 4 and selectorStart: # 上一道题选项提取完毕

            
            if selectorStart: # 提取选项
                # 一行包含多个选项，直接完成一道题目的提取
                if isMultipleSelector(text):
                    item["selector"] =  extractMultipleSelector(text)
                    if echo:
                        print(item)
                        sys.exit(0)
                    boxs.append(item)
                    # 重置状态
                    selectorStart = False
                    item = None
                    selector = []
                    continue
                else:
                    if extractMultipleSelector(text) is not None:
                        selector += extractMultipleSelector(text) 
                    else:
                        selector.append(text.strip())
                    
            else: # 提取题目表述
                # [ question, answer ] = extract_choice(text)
                    
                # if question is not None or answer is not None:
                question = remove_number(text) 
                answer = ""
                item = {
                    "question": question,
                    "answer": answer,
                }
                selector = [] # 重置选项
                selectorStart = True # 进行选项提取


            
    with open(curr_dir + "/result.json", "w", encoding="utf8") as f:
        json.dump(boxs, f, indent=4, ensure_ascii=False)

    # export_excel(boxs)
    write_excel_xls_append(boxs)

def test_extract_type():
    texts = [
        "一、单选题（20*1分）",
        "二、多选题",
        "   判断题 "
    ]
    results = ["单选题", "多选题", "判断题"]
    for t in texts:
        assert(extract_type(t) in results)
        
def text_extract_decide():
    texts = [
        "10.每种型号的红外测温仪都有自己特定的测温范围。（ √ ）",
        "3.“防”和“消”是不可分割的整体，“消”是“防”的先决条件，“消”必须与“防”紧密结合，“防”与“消”是实现消防安全的两种必要手段，两者互相联系，互相渗透，相辅相成，缺一不可。（ × ）",
        "24."
    ]
    
    for t in texts:
        result = extract_decide(t)
        print(result)
        
def text_is_multiple_selector():
    texts = [
        "A安全第一  B减少火灾危害   C生命至上   D以人为本",
        "A设置防火卷帘或防火幕等简易防火分隔物的上部	",
        "A SS50  B SS65  C SS100  D SS150",
        "A 200  B 50  C 150  D 100",
        "A12  B 24  C 36  D 48",
        "A地上式室外消火栓        B地下式室外消火栓"
        
    ] 
    for t in texts:
        result = isMultipleSelector(t)
        print(result)
        print(extractMultipleSelector(t))

    print(extract_choice("职业需具备的特征有（  ABCDE   ）"))
    print(extract_choice("12.灯具保养前，应确保系统保持主电源供电（ B ）h以上。"))
    
def text_new_selector():
    texts = [
        # "(A)每年(B) 每半年(C) 每季度(D) 每月",
        # "(A)三万元以上三十万元以下",
        # "(A)贯彻执行消防法规，保障单位消防安全符合规定，掌握本单位的消防安全情况;",
        # "(B)将消防工作与本单位的生产、科研、经营、管理等活动统筹安排，批准实施年度消防工作计划;为本单位的消防安全提供必要的经费和组织保障",
        # "(C)确定逐级消防安全责任，批准实施消防安全制度和保障消防安全的操作规程",
        # "(D)组织防火检查,督促落实火灾隐患整改，及时处理涉及消防安全的重大问题;",
        # "(E)组织制定符合本单位实际的灭火和应急疏散预案，并实施演练",
        # "(A)对火灾隐患不及时消除",
        # "(B)不按照国家有关规定配置消防设施和器材",
        # "(C)员工未进行消防安全培训",
        # "(D)未制订消防安全制度",
        # "(E) 占用、堵塞、封闭疏散通道、安全出口或者其他妨碍安全疏散行为的",
        # "消防管理法规(B)消防法规(C)部门规章(D)中华人民共和国消防法"
        "(A) 119报警电话是免费的", 
        "(A)外B)内(C)左(D)右"
    ]
    for t in texts:
        print(extractMultipleSelector(t))

    print(isNumberStart("20.单位违反 《消防法》规定，有下列哪些行为____的，责令改正并处五千以上五万以下罚款。"))
    print(isNumberStart("(A)保安部门(B)专兼职消防管理人员(C)法定代表人(D)消防安全管理"))
    print(isNumberStart("(A)可以给予支援"))

def test_main():
    # test_extract_type()
    # text_extract_decide()
    # text_is_multiple_selector()
    text_new_selector()


if __name__ == "__main__":
    test_main()
    
    analyze_document()