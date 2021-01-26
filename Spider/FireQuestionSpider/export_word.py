#!/usr/bin/env python3
#-*- coding: utf-8 -*-

import os
from pathlib import Path
import sys
import json

from docx import Document

from mysql import MySQL
import settings
from utils import *

# 测试MySQL
db = MySQL(settings.DATABASE_HOST, settings.DATABASE_USERNAME, settings.DATABASE_PASSWORD, settings.DATABASE_NAME)

# 应用文件目录
curr_dir = os.path.dirname(os.path.realpath(__file__))
word_dir = os.path.join(curr_dir, "word")

def init_env():
    """ 
    初始化环境
    """ 
    # 创建缓存目录
    if not os.path.exists(word_dir):
        os.mkdir(word_dir)
    os.chdir(word_dir)
    
def export_project(project_id):
    """ 
    导出单个项目为 word
    """ 
    # 查询项目名称
    sql = "SELECT name from `tk_chapter` WHERE id = %s"
    project_name = db.select_one(sql, (project_id))["name"]
    
    # 查询项目的章节
    sql = "SELECT id, name from `tk_chapter` WHERE pid = %s"
    chapters = db.select_all(sql, (project_id))
    
    # 查询章节对应的知识点
    for i, item in enumerate(chapters):
        sql = "SELECT id, name from `tk_chapter` WHERE pid = %s"
        knows = db.select_all(sql, (item["id"]))
        
        # 查询知识点对应的题目
        for j, know in enumerate(knows):
            sql = "SELECT `title`,`content`,`select`,`answer` from `tk_questions` WHERE chapter_id = %s"
            questions = db.select_all(sql, (know["id"]))
            knows[j]["questions"] = questions

        chapters[i]["child"] = knows

    # with open("questions.json", "w", encoding="utf8") as f:
    #     f.write(json.dumps(chapters))       

    document = Document() 
    document.add_heading(project_name, 0)
    
    for chapter in chapters:
        document.add_heading(chapter["name"], 1)
        for know in chapter["child"]: 
            document.add_heading(know["name"], 2)
            for question in know["questions"]:
                document.add_paragraph(question["title"])
                document.add_paragraph(question["content"])
                document.add_paragraph(question["select"])
                document.add_paragraph(question["answer"])

    filename = project_name.strip().replace("/", " ")
    document.save(filename + ".docx")

def export_all():
    sql = "SELECT id, name from `tk_chapter` WHERE pid = 0"
    projects = db.select_all(sql)
    for item in projects:
        export_project(item["id"])
        # os.rename("{0}.docx".format(item["id"]), "{0}.docx".format(item["name"].lrt))
    
if __name__ == "__main__":
    init_env()
    
    export_all()
    