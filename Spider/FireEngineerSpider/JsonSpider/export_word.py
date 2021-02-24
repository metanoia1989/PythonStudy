#!/usr/bin/env python3
#-*- coding: utf-8 -*-

import os
from pathlib import Path
import sys
import json

from docx import Document

from mysql import MySQL

# 测试MySQL
db = MySQL("localhost", "root", "root", "fire_engineer")

# 应用文件目录
curr_dir = os.path.dirname(os.path.realpath(__file__))
word_dir = os.path.join(curr_dir, "word")

def init_env():
    """ 
    初始化环境
    """ 
    # 创建缓存目录
    if not os.path.exists(word_dir):
        os.mkdir(word_dir)
    os.chdir(word_dir)

def br2n(content):
    """
    <br> 替换为 \n 换行符
    """
    return content.replace("<br>", "\n")  \
        .replace("&emsp;", "  ") \
        .replace("<br />", "\n")
    
def export_project(project_id):
    """ 
    导出单个项目为 word
    """ 
    # 查询项目名称
    sql = "SELECT name from `fire_chapter` WHERE id = %s"
    project_name = db.select_one(sql, (project_id))["name"]
    
    # 查询项目的章节
    sql = "SELECT id, name from `fire_chapter` WHERE pid = %s"
    chapters = db.select_all(sql, (project_id))
    
    # 查询章节对应的题目
    for i, chapter in enumerate(chapters):
        sql = "SELECT * FROM `fire_questions` WHERE chapter_id = %s"
        questions = db.select_all(sql, (chapter["id"]))
        chapters[i]["questions"] = questions

    project_dir = os.path.join(word_dir, project_name)
    if not os.path.exists(project_dir):
        os.mkdir(project_dir)
    os.chdir(project_dir) 

    
    for chapter in chapters:
        document = Document() 
        document.add_heading(chapter["name"], 0)
        for question in chapter["questions"]:
            title = document.add_paragraph()
            title.add_run(question["title"]).bold = True
            document.add_paragraph(br2n(question["content"]))
            document.add_paragraph(br2n(question["select"]))
            document.add_paragraph("答案：" + question["answer"])
            document.add_paragraph("解析：" + br2n(question["analyze"]))
            document.add_paragraph("")

        filename = chapter["name"].strip().replace("/", " ")
        document.save(filename + ".docx")
        

def export_all():
    sql = "SELECT id, name from `fire_chapter` WHERE pid = 0"
    projects = db.select_all(sql)
    for item in projects:
        export_project(item["id"])
        # os.rename("{0}.docx".format(item["id"]), "{0}.docx".format(item["name"].lrt))
    
if __name__ == "__main__":
    init_env()
    
    export_all()
    # export_project(1)
    